import sqlite3
import pandas
import docx


# ВАЖНО! В заданиие в примерах отсчетов присутствуют года 2006 и 2020, однако 2006 года в принципе нет в базе данных,
# а 2020 не подходит по заданные условия в запросе. По этой причине в файлы отсчетов я их не включал, но добавить
# пустые ячейки можно изменив пару строк кода

# Были мысли немного по-другому организовать класс Report, но я остановился на данном этом варианте.
# Рассмтаривал добавление метода get_excel, но мне показалось, что в рамках данной задачи это не является необходимым

class Report:
    def __init__(self):
        # При инициализации класса подключаемся к базе данных, создаем dataframe на основе запроса,
        # создаем все необходимые в последующем переменные

        self._db = sqlite3.connect('test.db')

        self._query = pandas.read_sql_query('''
        SELECT * FROM testidprod
        WHERE partner IS NULL AND state IS NULL AND bs=0 AND (factor=1 OR factor=2)
        ''', self._db)

        self._report = pandas.DataFrame(self._query)
        self._df_factor_6 = pandas.DataFrame()

        self._report = self._report[['factor', 'year', 'res']].groupby(['factor', 'year']).sum()
        self._report = self._report.rename(columns={'factor': 'Factor', 'year': 'Year', 'res': 'World'}).transpose()

        self._factor_1, self._factor_2, self._factor_6, self._years, self._indexes = [], [], [], [], []

        # Эти две переменные используются для отслеживания использованных на объекте методов, чтобы предотвратить
        # ошибки при удалении/добавлении фактора 6
        self._add_was_run = False
        self._delete_was_run = False

    def get_report(self):
        # Метод возвращает нынешний отсчет
        return self._report

    def add_factor_6(self):
        # Метод добавляет фактор 6, если он не существует, далее экспортирует self._report в excel

        for col in self._report.columns:
            if col[0] == 1:
                self._factor_1.append(col)
            if col[0] == 2:
                self._factor_2.append(col)

        if not self._factor_6:
            for i in range(len(self._factor_2)):
                self._years.append(self._factor_2[i][1])
                if (1, self._factor_2[i][1]) in self._factor_1:
                    self._factor_6.append(self._report[self._factor_2[i]]['World']
                                          / self._report[(1, self._factor_2[i][1])]['World'])
                else:
                    self._factor_6.append('NaN')

            for i in range(len(self._factor_6)):
                self._indexes.append(6)

            self._df_factor_6 = (pandas
                                 .DataFrame({'World': self._factor_6},
                                            index=[self._indexes, self._years]).transpose())

        if self._delete_was_run is True or (self._delete_was_run is False and self._add_was_run is False):
            self._report = pandas.concat([self._report, self._df_factor_6], axis=1)
            self._delete_was_run = False
            self._add_was_run = True

        self._report.to_excel('report.xlsx')

        return self._report

    def delete_factor_6(self):
        # Метод удаляет фактор 6, если он существует, далее экспортирует self._report в excel

        if self._add_was_run:
            self._report = self._report.transpose()
            for element in self._factor_2:
                self._report['World'] = self._report['World'].drop((6, element[1]), axis=0)
            self._report = self._report.dropna().transpose()
            self._report.to_excel('report.xlsx')

            self._delete_was_run = True
            self._add_was_run = False

        return self._report

    def get_docx(self, _start=2007, _end=2019, **kwargs):
        # Метод может быть использован только после метода add_factor_6
        def get_cagr(values, start, end):

            values = [values[x] for x in range(len(values)) if self._years.index(start) <= x <= self._years.index(end)]

            cagr = round((((values[len(values) - 1] / values[0]) ** (1 / abs(end - start)) - 1) * 100), 2)
            grew = 'grew'

            if cagr < 0:
                grew = 'decreased'

            return f'Factor 6 {grew} by avg {abs(cagr)}% every year' \
                   f' from {start} to {end}.'

        if kwargs.get('start') and kwargs.get('start') >= self._years[0]:
            _start = kwargs.get('start')
        if kwargs.get('end') and kwargs.get('end') <= self._years[-1]:
            _end = kwargs.get('end')

        _values = self._factor_6

        if self._add_was_run is True and self._delete_was_run is False:
            _word_report = self._df_factor_6.transpose()
            _doc = docx.Document()

            _docx_table = _doc.add_table(len(self._factor_6) + 1, 3)
            _docx_table.style = 'Table Grid'

            _docx_table.cell(0, 0).text = 'Factor'
            _docx_table.cell(0, 1).text = 'Year'
            _docx_table.cell(0, 2).text = 'World Value'

            for i in range(1, len(self._factor_6) + 1):
                for j in range(3):
                    _docx_table.cell(i, 0).text = str(6)
                    _docx_table.cell(i, 1).text = str(round(_word_report['World'].index[i - 1][1], 2))
                    _docx_table.cell(i, 2).text = str(round(_word_report['World'].values[i - 1], 2))

            _docx_table.columns[0].cells[1].merge(_docx_table.columns[0].cells[len(self._factor_6)])
            _docx_table.cell(1, 0).text = '6'
            _docx_table.cell(1, 0).vertical_alignment = 1

            _p = _doc.add_paragraph()
            _p.text = get_cagr(_values, _start, _end)

            _doc.save('report.docx')
        else:
            print('Фактор 6 не существует, создайте его при помощи метода add_factor_6')

        return 'success'


if __name__ == '__main__':
    report = Report()

    report.add_factor_6()
    report.get_docx()
